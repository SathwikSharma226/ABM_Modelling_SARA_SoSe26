"""Build the project documentation .docx file (Milestone 1)."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_mono_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def add_kv_table(rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
        for para in cells[0].paragraphs:
            for r in para.runs:
                r.bold = True
    return table


# =========================================================================
# Cover
# =========================================================================
title = doc.add_heading("Waste in the City — Agent-Based Model", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Milestone 1 — Design & Architecture Review")
run.bold = True
run.font.size = Pt(14)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    "Master's Module: Symbolic AI and Rule-based Agents (SARA)\n"
    "OTH Amberg-Weiden — Summer Semester 2026\n"
    "Supervisor: Prof. Dr. Tatyana Ivanovska\n"
    "Project specification: lecture5_1_SARA.pdf, slides 41–48"
)
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph()

add_heading("Milestone 1 — Meeting Briefing", level=1)
add_kv_table([
    ("Milestone", "M1 — Design, Architecture & Agent Specification"),
    ("Status", "Design complete; reference implementation scaffolded"),
    ("Purpose of meeting",
     "Present the proposed architecture, agent catalogue, map model and "
     "movement strategy; obtain supervisor feedback before locking the "
     "design and moving to full implementation and experimentation."),
    ("Scope of this document",
     "(1) Project architecture and file structure, "
     "(2) agent catalogue and rule sets, "
     "(3) planned map generation, "
     "(4) agent mapping and movement."),
    ("Out of scope for M1",
     "Final experimental results, statistical analysis of the slide-46 "
     "sweeps, project defence slides, and the written final report. "
     "These belong to Milestones 2 and 3."),
])

add_heading("Milestone Plan (Overview)", level=2)
add_para(
    "The project is structured into three milestones. This meeting "
    "concludes Milestone 1; the deliverables of Milestones 2 and 3 are "
    "listed here only to give context to the present scope."
)
add_bullet("M1 — Design & Architecture (this milestone): project "
           "structure, agent specification, map representation, "
           "movement and pathfinding strategy, configuration model.")
add_bullet("M2 — Implementation & Experiments: complete the Mesa "
           "integration, run the slide-46 experiment battery (cleaner "
           "strategies, tourist density, bin/cleaning capacity), and "
           "validate the creative extension (decaying heatmap).")
add_bullet("M3 — Evaluation & Defence: statistical evaluation of the "
           "metrics, written report, defence presentation.")

add_heading("Discussion Points for the Supervisor", level=2)
add_bullet("Is the chosen agent taxonomy (locals, tourists, cleaners, "
           "bins, transporters) sufficient to cover the slide-41–48 "
           "requirements, or should additional agent types be added?")
add_bullet("Are the proposed cleaning strategies (nearest_waste, "
           "random_patrol, fixed_route, heatmap) adequate as the four "
           "comparison conditions for the slide-46 experiment?")
add_bullet("Is the decaying-heatmap design an acceptable creative "
           "extension for the slide-47 requirement?")
add_bullet("Is the ASCII-based map representation acceptable, or "
           "should we move to a procedurally generated city for M2?")
add_bullet("Confirm the chosen metrics (GroundWaste, OverflowingBins, "
           "AvgBinFill, WasteCleaned, WasteDisposed, WasteIntoBins) "
           "cover the slide-46 research questions.")

doc.add_page_break()

# =========================================================================
# 1. Architecture
# =========================================================================
add_heading("1. Project Architecture and File Structure", level=1)

add_para(
    "The project is implemented in Python on top of the Mesa "
    "agent-based modelling framework. It follows a flat, modular "
    "architecture in which each file owns one concern: configuration, "
    "search algorithms, agent behaviour, world state, visualisation, "
    "experiments, and the CLI entry point. This separation keeps the "
    "code easy to read for the milestone reviews and easy to extend "
    "with additional strategies or layouts in Milestone 2."
)

add_heading("1.1 Folder Layout", level=2)
add_mono_block(
    "ABM_Modelling_SARA_SoSe26/\n"
    "├── README.md              # User-facing project overview\n"
    "├── requirements.txt       # Python dependencies\n"
    "├── config.py              # All tunable parameters (single source of truth)\n"
    "├── pathfinding.py         # BFS / multi-target BFS / A* graph search\n"
    "├── agents.py              # The five rule-based agent classes\n"
    "├── city_model.py          # Mesa Model: layout parsing, scheduler, DataCollector\n"
    "├── visualize.py           # Matplotlib snapshot, animation, metrics plots\n"
    "├── experiments.py         # Slide-46 comparison experiment battery\n"
    "├── run.py                 # Command-line entry point\n"
    "├── city_layouts/\n"
    "│   ├── README.md          # Legend for ASCII layout files\n"
    "│   └── default.txt        # Default 30×30 city map\n"
    "└── results/               # Experiment artefacts — populated in M2\n"
)

add_heading("1.2 File-by-File Overview", level=2)

files = [
    ("config.py",
     "Central configuration module. Every tunable lever of the model "
     "lives here so experiments can be reproduced and so reviewers "
     "can see all knobs in one place: grid size, agent population "
     "counts, waste-drop probabilities, bin and container capacities, "
     "the whitelist of cleaner strategies, the heatmap decay and "
     "increment values, the number of simulation steps, and the "
     "random seed."),
    ("pathfinding.py",
     "Self-contained graph-search utility module. Implements "
     "bfs_shortest_path (single-source/single-goal BFS), bfs_nearest "
     "(multi-target BFS used to find the closest bin or piece of "
     "waste), and astar_path (A* with Manhattan-distance heuristic). "
     "The graph is implicit: nodes are walkable grid cells, edges "
     "connect 4-connected neighbours (no diagonal movement)."),
    ("agents.py",
     "Defines the five mesa.Agent subclasses (LocalHumanAgent, "
     "TouristAgent, CleaningServiceAgent, DustBinAgent, "
     "DustTransporterAgent) together with a shared _MovingAgent base "
     "class that provides random_walk() and step_along_path() "
     "helpers. Every step() method is heavily commented to support "
     "the milestone reviews."),
    ("city_model.py",
     "Houses the WasteCityModel (mesa.Model). Responsible for parsing "
     "the ASCII layout, instantiating the static infrastructure "
     "(bins, containers, disposal points), spawning the requested mix "
     "of mobile agents, maintaining global state (ground waste "
     "dictionary and decaying heatmap), and configuring the Mesa "
     "DataCollector that samples time-series metrics on every tick."),
    ("visualize.py",
     "Pure Matplotlib visualisation layer. Renders the static cells "
     "as a coloured background and overlays bins (with current load), "
     "waste piles, and mobile agents. Provides draw_city() for static "
     "snapshots, animate() for live runs (with optional GIF export), "
     "and plot_metrics() for the time-series charts."),
    ("experiments.py",
     "Implements the experiment battery from slide 46: cleaner-"
     "strategy sweep, tourist-density sweep, and bin/cleaning-"
     "capacity sweep. Writes comparison plots and the raw metric CSVs "
     "into the results/ folder. Full execution is scheduled for "
     "Milestone 2."),
    ("run.py",
     "Friendly command-line entry point. Exposes flags to override "
     "the strategy, population mix, step count, animation behaviour, "
     "GIF/PNG output paths, and a single --experiments switch that "
     "runs the full battery."),
    ("city_layouts/default.txt",
     "The default ASCII map (30×30). Each character is one cell: '.' "
     "street, '#' building, 'P' public area, 'A' attraction, 'B' "
     "bin, 'C' container, 'D' disposal point."),
    ("results/",
     "Output folder populated by --experiments. Will hold the "
     "comparison PNGs (compare_strategies_*, tourists_*, "
     "bin_density_*) and the raw per-run CSV time series produced "
     "during Milestone 2."),
]
for name, desc in files:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(name)
    run.bold = True
    run.font.name = "Consolas"
    p.add_run(" — " + desc)

add_heading("1.3 Module Dependencies", level=2)
add_para(
    "The dependency graph is intentionally shallow and acyclic. "
    "config.py has no internal imports. pathfinding.py is standalone. "
    "agents.py depends on pathfinding.py and config.py. city_model.py "
    "depends on agents.py and config.py. visualize.py and "
    "experiments.py depend on city_model.py. run.py is the only "
    "module that wires everything together for the user."
)

add_heading("1.4 Milestone 1 Status", level=2)
add_bullet("Folder structure, configuration module and dependency "
           "graph are fixed.")
add_bullet("Pathfinding utilities (BFS, multi-target BFS, A*) are "
           "implemented and unit-testable.")
add_bullet("All five agent classes are scaffolded with their step() "
           "rules defined as described in §2.")
add_bullet("The Mesa model parses the ASCII layout, spawns agents and "
           "registers the DataCollector reporters.")
add_bullet("Visualisation and the experiment battery are stubbed; "
           "full execution and tuning are tasks for Milestone 2.")

# =========================================================================
# 2. Agents and Rules
# =========================================================================
add_heading("2. List of Agents and Their Rules", level=1)

add_para(
    "Five rule-based agent types are proposed for the model, "
    "mirroring the lecture specification (slides 41–48). Two are "
    "static infrastructure, three are mobile. All inherit from "
    "mesa.Agent and implement a step() method that fires once per "
    "simulation tick under Mesa's RandomActivation scheduler."
)

add_heading("2.1 LocalHumanAgent", level=2)
add_para("Role: resident that follows a daily home ↔ work commute "
         "pattern and occasionally produces waste.")
add_para("Rules:", bold=True)
add_bullet("Spawns at its home cell; remembers a (home, work) pair "
           "selected from the walkable cells.")
add_bullet("On every step: if it has arrived at its current target, "
           "the target is toggled between home and work; the cached "
           "BFS path is then invalidated.")
add_bullet("Movement: BFS computes the shortest path on the "
           "4-connected walkable graph; the agent advances one cell "
           "per step along the cached path. If the next cell is "
           "blocked, it falls back to a random walk.")
add_bullet("Waste production: with probability LOCAL_WASTE_PROB "
           "(default 0.05) per step, it considers dropping waste.")
add_bullet("Bin preference: searches for the nearest non-full bin "
           "via multi-target BFS. If a bin is reachable within "
           "HUMAN_BIN_SEARCH_RADIUS (default 5 cells), the waste goes "
           "into the bin and the waste_into_bins counter is "
           "incremented; otherwise the waste is dropped on the "
           "ground at the current cell.")

add_heading("2.2 TouristAgent", level=2)
add_para("Role: visitor that drifts toward attractions and litters "
         "more freely than locals.")
add_para("Rules:", bold=True)
add_bullet("On creation, picks a random 'A' attraction cell as its "
           "target (falls back to a random walkable cell if no "
           "attraction exists).")
add_bullet("With a 5% probability per step, or upon arrival, picks a "
           "new attraction as the next target (less predictable than "
           "locals).")
add_bullet("Movement: BFS toward the current attraction target, but "
           "with a 30% chance per step the tourist meanders via a "
           "random walk instead of following the planned path.")
add_bullet("Waste production: with probability TOURIST_WASTE_PROB "
           "(default 0.15) the tourist drops waste — always on the "
           "ground, never into a bin (modelling lower waste "
           "awareness).")

add_heading("2.3 CleaningServiceAgent", level=2)
add_para("Role: street cleaner that picks up ground waste according "
         "to a rule-based strategy chosen at instantiation time.")
add_para("Rules:", bold=True)
add_bullet("Step-priority rule: if the cleaner is already standing "
           "on a ground-waste cell, it removes the waste, updates "
           "cleaned_units and the waste_cleaned counter, and ends "
           "the step.")
add_bullet("Otherwise, the cleaner computes a movement plan based on "
           "its strategy and advances one cell along it; if no plan "
           "is available it falls back to a random walk.")
add_para("Strategies (whitelisted in config.CLEANER_STRATEGIES):",
         bold=True)
add_bullet("nearest_waste — multi-target BFS to the closest known "
           "ground-waste cell.")
add_bullet("random_patrol — no global plan; the cleaner wanders the "
           "street network step by step.")
add_bullet("fixed_route — cycles through a preset list of waypoints "
           "(bin cells by default), advancing the route index on "
           "arrival.")
add_bullet("heatmap (creative extension) — BFS toward the cell with "
           "the highest decaying-heatmap score returned by "
           "model.heatmap_argmax().")

add_heading("2.4 DustBinAgent", level=2)
add_para("Role: static infrastructure that stores waste up to a "
         "capacity. Models both small 'bin' (B, capacity 10) and "
         "large 'container' (C, capacity 50) cells.")
add_para("Rules:", bold=True)
add_bullet("Exposes is_full and fill_ratio properties used by humans "
           "and transporters to make decisions.")
add_bullet("add_waste(amount) accepts as much waste as free capacity "
           "allows; any leftover counts as an overflow event and is "
           "tracked in overflow_count.")
add_bullet("empty() removes and returns the current load (used by "
           "the transporter).")
add_bullet("step() is a no-op — bins are passive but are scheduled "
           "so Mesa treats them uniformly.")

add_heading("2.5 DustTransporterAgent", level=2)
add_para("Role: truck-style agent that empties sufficiently-full "
         "bins and delivers the waste to a disposal point ('D').")
add_para("Rules:", bold=True)
add_bullet("Maintains an onboard cargo counter and a current target "
           "bin.")
add_bullet("If standing on the depot with cargo > 0, the cargo is "
           "dumped: waste_disposed is incremented and the waste is "
           "permanently removed from the simulation.")
add_bullet("If standing on the target bin, the bin is emptied into "
           "the cargo, bins_emptied is incremented, and a path back "
           "to the depot is planned.")
add_bullet("Target selection: picks the closest bin whose fill_ratio "
           "is at or above TRANSPORTER_FULL_THRESHOLD (default 0.7) "
           "via multi-target BFS — avoiding pointless trips to "
           "nearly-empty bins.")
add_bullet("Movement: advances one cell per step along the cached "
           "BFS path; replans when the path is consumed or "
           "invalidated.")

# =========================================================================
# 3. Map generation
# =========================================================================
add_heading("3. How the Map Is Planned to Be Generated", level=1)

add_para(
    "For Milestone 1, the city is not procedurally generated at "
    "runtime. Instead, the map is authored once as a plain-text ASCII "
    "file and parsed by the model on construction. This design keeps "
    "the world deterministic, makes the city structure visible at a "
    "glance, and lets us hand-craft alternative layouts for "
    "experiments without touching the code. Procedural generation is "
    "recorded as an optional extension for later milestones."
)

add_heading("3.1 ASCII Layout Format", level=2)
add_para(
    "Every character in the layout file represents one grid cell. "
    "The legend mirrors the constants documented in config.py:"
)
add_mono_block(
    "Symbol  Meaning\n"
    "------  -------------------------------------------------\n"
    "  .     Street / walkable path\n"
    "  #     Wall / building (blocks movement)\n"
    "  P     Public area (walkable; tends to collect waste)\n"
    "  A     Attraction (tourists gravitate here)\n"
    "  B     Dust bin (small static infrastructure)\n"
    "  C     Dust container (large static infrastructure)\n"
    "  D     Disposal point (transporter terminus)\n"
)

add_heading("3.2 Default 30×30 Map", level=2)
add_para(
    "city_layouts/default.txt provides a sample downtown: a central "
    "plaza made of 'P' public cells with two 'A' attractions, "
    "surrounded by a regular grid of building blocks and street "
    "corridors. Bins ('B') are scattered along the main streets, two "
    "larger containers ('C') sit near the plaza, and a single "
    "disposal point ('D') is placed at the upper-left edge of the "
    "city."
)
add_mono_block(
    "D....B.....B.....B.....B......\n"
    ".#####.#####.#####.#####.#####\n"
    ".#####.#####.#####.#####.#####\n"
    ".#####.#####.#####.#####.#####\n"
    ".#####.#####.#####.#####.#####\n"
    ".#####.#####.#####.#####.#####\n"
    ".....B.................B......\n"
    ".#####.#####.#####.#####.#####\n"
    ".....B......C..........B......\n"
    ".#####.#####.PPPP#.#####.#####\n"
    ".#####.#####.PAPP#.#####.#####\n"
    ".#####.#####.PPAP#.#####.#####\n"
    ".#####.#####.PPPP#.#####.#####\n"
    ".....B............C....B......\n"
    "...   (... continues for 30 rows ...)   ...\n"
)

add_heading("3.3 Parsing Pipeline", level=2)
add_para(
    "When a WasteCityModel is constructed, the map is materialised "
    "through the following deterministic pipeline:"
)
add_bullet("load_layout(path) opens the file with UTF-8 encoding, "
           "strips newlines, skips empty lines, and right-pads "
           "shorter rows with '#' so the grid becomes a strict "
           "rectangle.")
add_bullet("The returned (rows, width, height) tuple is stored on "
           "the model as model.cell_types, indexed as "
           "cell_types[y][x].")
add_bullet("A mesa.space.MultiGrid of the parsed dimensions is "
           "created with torus=False so the city has hard borders.")
add_bullet("_scan_layout walks the character grid once. For each "
           "cell it (a) records 'A' cells in model.attractions, "
           "(b) records every non-building cell in "
           "model.walkable_cells, (c) instantiates a DustBinAgent on "
           "each 'B' (small) and 'C' (large) cell, and (d) records "
           "each 'D' cell as a disposal point. If no 'D' is present, "
           "(0, 0) is used as a fallback so the transporter always "
           "has a valid depot.")
add_bullet("Mobile agents are then spawned: locals receive random "
           "(home, work) pairs chosen from walkable_cells; tourists "
           "spawn on random walkable cells; cleaners spawn on random "
           "walkable cells with a fixed route synthesised from the "
           "first half of the bins; transporters spawn directly on "
           "the first disposal point.")

add_heading("3.4 Adding New Layouts", level=2)
add_para(
    "New maps are added by dropping a .txt file into city_layouts/ "
    "that uses only the legend characters above. The map is then "
    "activated either by setting config.DEFAULT_LAYOUT_FILE or by "
    "passing layout_file=... to the WasteCityModel constructor. The "
    "only authoring constraint is that at least one 'D' cell must be "
    "reachable from the bins; otherwise the transporter cannot "
    "dispose of waste."
)

# =========================================================================
# 4. Mapping and movement
# =========================================================================
add_heading("4. Possible Mapping and Movement of Agents", level=1)

add_heading("4.1 The Movement Graph", level=2)
add_para(
    "The world is modelled as a 4-connected grid graph. Nodes are "
    "grid cells; an edge connects two cells if (a) they share a "
    "north/south/east/west boundary and (b) both cells are walkable. "
    "Diagonal moves are disallowed so paths follow streets "
    "naturally. Walkability is decided by the make_walkable(model) "
    "predicate factory in agents.py: building cells ('#') and "
    "out-of-bounds coordinates are rejected; everything else "
    "(streets, public areas, attractions, bins, containers and "
    "disposal points) is walkable."
)

add_heading("4.2 Pathfinding Algorithms", level=2)
add_para(
    "Three graph-search algorithms in pathfinding.py back every "
    "movement decision, fulfilling the lecture requirement to use "
    "graph search where it makes sense:"
)
add_bullet("bfs_shortest_path(start, goal, walkable) — breadth-first "
           "search returning the shortest path as a list of cells "
           "(including start and goal); returns None when the goal "
           "is unreachable. Used whenever an agent knows a specific "
           "target.")
add_bullet("bfs_nearest(start, targets, walkable) — multi-target BFS "
           "that returns the closest reachable target and the path "
           "to it in a single traversal. Used by humans to find the "
           "nearest bin and by cleaners/transporters to find the "
           "closest waste pile or sufficiently-full bin.")
add_bullet("astar_path(start, goal, walkable) — A* with a Manhattan "
           "distance heuristic, admissible on a 4-connected unit-"
           "cost grid. Available for goal-directed search where the "
           "heuristic offers a clear speed-up.")

add_heading("4.3 Per-Agent Movement Plans", level=2)

add_para("LocalHumanAgent:", bold=True)
add_bullet("Targets oscillate between home and work each time the "
           "agent arrives; one BFS path is computed per target and "
           "cached.")
add_bullet("The agent consumes one cell of the cached path per "
           "step; the consumed cell is dropped so the next call "
           "resumes from the current position.")
add_bullet("If the planned next cell becomes non-walkable (defensive "
           "re-check), the agent falls back to a single random walk "
           "step.")

add_para("TouristAgent:", bold=True)
add_bullet("Targets are attraction cells, replaced randomly with 5% "
           "probability per step or whenever the tourist arrives.")
add_bullet("Movement alternates between planned BFS steps (70% of "
           "steps) and pure random walks (30% of steps), producing "
           "a characteristic meandering trajectory.")

add_para("CleaningServiceAgent:", bold=True)
add_bullet("nearest_waste: bfs_nearest from the cleaner's cell to "
           "the set of ground-waste positions; the cleaner steps "
           "along the shortest route to the closest pile.")
add_bullet("random_patrol: no plan is produced; the agent always "
           "falls back to a random walk on the street network.")
add_bullet("fixed_route: bfs_shortest_path to the current waypoint "
           "in a preset list (the first half of the bin cells); "
           "when the waypoint is reached, the route index advances "
           "cyclically.")
add_bullet("heatmap: bfs_shortest_path to model.heatmap_argmax(), "
           "the cell with the highest decaying-heatmap value; this "
           "lets the cleaner pre-position around chronic hotspots "
           "even when individual waste pieces are not yet visible "
           "to the greedy nearest-waste heuristic.")

add_para("DustTransporterAgent:", bold=True)
add_bullet("State machine with three phases: select-bin → "
           "travel-to-bin → travel-to-depot.")
add_bullet("Bin selection uses bfs_nearest restricted to bins whose "
           "fill_ratio ≥ TRANSPORTER_FULL_THRESHOLD (0.7), avoiding "
           "trips to nearly-empty bins.")
add_bullet("After emptying a bin, a BFS path back to the depot is "
           "computed and consumed step by step; on arrival at the "
           "depot the cargo is permanently disposed and the cycle "
           "restarts.")

add_para("DustBinAgent:", bold=True)
add_bullet("Does not move. Bins are fixed at the (x, y) coordinate "
           "of their source 'B' or 'C' character in the ASCII map "
           "for the entire run.")

add_heading("4.4 Decaying Heatmap (Creative Extension)", level=2)
add_para(
    "On top of the reactive movement above, the model maintains a "
    "2-D heatmap of waste observations. Each time waste lands on "
    "the ground the heatmap value at that cell is incremented by "
    "HEATMAP_INCREMENT (default 1.0). At the start of every model "
    "step the whole heatmap is multiplied by HEATMAP_DECAY (default "
    "0.97), so transient incidents fade exponentially while "
    "persistent hotspots stay strong. The heatmap cleaner strategy "
    "plans its movement against this memory, giving cleaners "
    "predictive — not just reactive — behaviour. The full "
    "quantitative comparison against the other three strategies is "
    "scheduled for Milestone 2."
)

add_heading("4.5 Step Order and Determinism", level=2)
add_para(
    "WasteCityModel.step() executes in a fixed order on every tick: "
    "(1) DataCollector samples all model reporters so step 0 "
    "reflects the initial state; (2) the heatmap is decayed in "
    "place; (3) Mesa's RandomActivation schedules every registered "
    "agent exactly once in random order. Reproducibility is "
    "achieved by seeding the model's random number generator via "
    "config.RANDOM_SEED (default 42)."
)

# =========================================================================
# Next steps & sign-off
# =========================================================================
doc.add_page_break()
add_heading("Next Steps After Milestone 1", level=1)
add_bullet("Incorporate supervisor feedback from this meeting into "
           "the agent rules and the cleaner-strategy whitelist.")
add_bullet("Lock the configuration defaults (population mix, "
           "probabilities, capacities) used for the slide-46 sweeps.")
add_bullet("Run the full experiment battery (cleaner strategy, "
           "tourist density, bin/cleaning capacity) and collect "
           "time-series metrics into results/.")
add_bullet("Quantify the impact of the decaying-heatmap creative "
           "extension against the three baseline strategies.")
add_bullet("Prepare the Milestone 2 deliverable: experiment plots, "
           "CSV artefacts and a short results write-up.")

add_heading("Sign-off", level=1)
add_kv_table([
    ("Student", "____________________________"),
    ("Supervisor", "Prof. Dr. Tatyana Ivanovska"),
    ("Date of Milestone 1 meeting", "____________________________"),
    ("Outcome",
     "[ ] Approved as-is   [ ] Approved with changes   [ ] Revisit"),
    ("Action items / changes requested",
     "____________________________________________________________"),
])

out = "Waste_in_the_City_Project_Documentation.docx"
doc.save(out)
print(f"Wrote {out}")
